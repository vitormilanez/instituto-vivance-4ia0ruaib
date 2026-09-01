#!/usr/bin/env python3
"""Convert the Vivance Markdown source into a Google Docs-targeted DOCX.

The document uses the `google_docs_default` preset from the bundled documents
skill. The title is deliberately a plain paragraph, not Word's Title style.
"""

from __future__ import annotations

import argparse
import math
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor, Twips
from PIL import Image, ImageDraw, ImageFont


DOCUMENT_SKILL_ROOT = Path(
    "/Users/vitormilanez/.cache/codex-runtimes/codex-primary-runtime/"
    "plugins/openai-primary-runtime/plugins/documents/skills/documents"
)
sys.path.insert(0, str(DOCUMENT_SKILL_ROOT / "scripts"))
from table_geometry import apply_table_geometry, column_widths_from_weights  # noqa: E402


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
FONT = "Arial"
MONO_FONT = "Courier New"
INK = "000000"
MUTED = "555555"
BORDER = "DADCE0"
CODE_FILL = "F8F9FA"
CONTENT_WIDTH_DXA = 9360


def set_attr(element, name: str, value: str) -> None:
    element.set(qn(name), value)


def ensure_child(parent, tag: str):
    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        parent.append(child)
    return child


def set_run_font(run, name: str = FONT, size: float | None = None, color: str = INK) -> None:
    run.font.name = name
    if run._element.rPr is None:
        run._element.get_or_add_rPr()
    fonts = run._element.rPr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        run._element.rPr.insert(0, fonts)
    fonts.set(qn("w:ascii"), name)
    fonts.set(qn("w:hAnsi"), name)
    fonts.set(qn("w:eastAsia"), name)
    run.font.color.rgb = RGBColor.from_string(color)
    if size is not None:
        run.font.size = Pt(size)


def set_style_font(style, name: str, size: float, color: str, bold: bool | None = None) -> None:
    style.font.name = name
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), name)


def set_paragraph_spacing(style, *, before: float, after: float, line: float) -> None:
    fmt = style.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    set_style_font(normal, FONT, 11, INK, False)
    set_paragraph_spacing(normal, before=0, after=8, line=1.15)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    normal.paragraph_format.widow_control = True

    heading_tokens = {
        "Heading 1": (20, INK, False, 20, 6),
        "Heading 2": (16, INK, False, 18, 6),
        "Heading 3": (14, "434343", False, 16, 4),
        "Heading 4": (12, "434343", True, 12, 3),
    }
    for name, (size, color, bold, before, after) in heading_tokens.items():
        style = styles[name]
        set_style_font(style, FONT, size, color, bold)
        set_paragraph_spacing(style, before=before, after=after, line=1.15)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True
        style.paragraph_format.page_break_before = False

    if "Vivance Metadata" not in styles:
        metadata = styles.add_style("Vivance Metadata", WD_STYLE_TYPE.PARAGRAPH)
    else:
        metadata = styles["Vivance Metadata"]
    set_style_font(metadata, FONT, 10.5, MUTED, False)
    set_paragraph_spacing(metadata, before=0, after=3, line=1.15)
    metadata.paragraph_format.left_indent = Inches(0.25)

    if "Vivance Code Block" not in styles:
        code_style = styles.add_style("Vivance Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code_style = styles["Vivance Code Block"]
    set_style_font(code_style, MONO_FONT, 9, "202124", False)
    set_paragraph_spacing(code_style, before=4, after=6, line=1.08)
    code_style.paragraph_format.left_indent = Inches(0.2)
    code_style.paragraph_format.right_indent = Inches(0.2)
    code_style.paragraph_format.keep_together = True

    if "Vivance Table Text" not in styles:
        table_text = styles.add_style("Vivance Table Text", WD_STYLE_TYPE.PARAGRAPH)
    else:
        table_text = styles["Vivance Table Text"]
    set_style_font(table_text, FONT, 9.5, INK, False)
    set_paragraph_spacing(table_text, before=0, after=0, line=1.05)

    core = doc.core_properties
    core.title = "Instituto Vivance — Contexto do Produto, Escopo do MVP e Casos de Uso"
    core.subject = "Documento-base para discovery, produto, design e desenvolvimento"
    core.author = "Instituto Vivance"
    core.keywords = "Instituto Vivance, MVP, cuidado longitudinal, IA clínica, casos de uso"
    core.comments = "Conversão integral do documento Markdown para Google Docs."


def add_hyperlink(paragraph, text: str, url: str, *, bold: bool = False, italic: bool = False) -> None:
    relationship_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), FONT)
    fonts.set(qn("w:hAnsi"), FONT)
    rpr.append(fonts)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1155CC")
    rpr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(underline)
    if bold:
        rpr.append(OxmlElement("w:b"))
    if italic:
        rpr.append(OxmlElement("w:i"))
    run.append(rpr)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


INLINE_RE = re.compile(
    r"(\*\*[^*]+?\*\*|(?<!\*)\*[^*]+?\*(?!\*)|`[^`]+?`|\[[^\]]+?\]\([^)]+?\))"
)


def add_inline(paragraph, text: str, *, bold: bool = False, italic: bool = False) -> None:
    cursor = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor : match.start()])
            set_run_font(run)
            run.bold = bold
            run.italic = italic
        token = match.group(0)
        if token.startswith("**"):
            add_inline(paragraph, token[2:-2], bold=True, italic=italic)
        elif token.startswith("*"):
            add_inline(paragraph, token[1:-1], bold=bold, italic=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, MONO_FONT, 9.5, "202124")
            run.bold = bold
            run.italic = italic
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), "F1F3F4")
            run._element.get_or_add_rPr().append(shading)
        elif token.startswith("["):
            link_match = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token)
            if link_match:
                add_hyperlink(
                    paragraph,
                    link_match.group(1),
                    link_match.group(2),
                    bold=bold,
                    italic=italic,
                )
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_run_font(run)
        run.bold = bold
        run.italic = italic


def set_paragraph_shading(paragraph, fill: str) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    shd = ensure_child(ppr, "w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def add_title(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    set_run_font(run, FONT, 26, INK)
    run.bold = False


def next_free_id(elements, attr_name: str) -> int:
    values = []
    for element in elements:
        value = element.get(qn(attr_name))
        if value is not None:
            values.append(int(value))
    return max(values, default=0) + 1


def create_numbering(doc: Document, kind: str, start_value: int = 1) -> int:
    numbering = doc.part.numbering_part.element
    abstracts = numbering.findall(qn("w:abstractNum"))
    nums = numbering.findall(qn("w:num"))
    abstract_id = next_free_id(abstracts, "w:abstractNumId")
    num_id = next_free_id(nums, "w:numId")

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "multilevel")
    abstract.append(multi)

    for level in range(9):
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(level))
        start = OxmlElement("w:start")
        start.set(qn("w:val"), str(start_value if level == 0 else 1))
        lvl.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), "decimal" if kind == "decimal" else "bullet")
        lvl.append(num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        if kind == "decimal":
            marker = f"%{level + 1}."
        elif kind == "checkbox":
            marker = "☐"
        else:
            marker = ["●", "○", "■"][level % 3]
        lvl_text.set(qn("w:val"), marker)
        lvl.append(lvl_text)
        lvl_jc = OxmlElement("w:lvlJc")
        lvl_jc.set(qn("w:val"), "left")
        lvl.append(lvl_jc)

        marker_pos = 360 + level * 360
        text_indent = 720 + level * 360
        ppr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), str(text_indent))
        tabs.append(tab)
        ppr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(text_indent))
        ind.set(qn("w:hanging"), str(text_indent - marker_pos))
        ppr.append(ind)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:before"), "0")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "276")
        spacing.set(qn("w:lineRule"), "auto")
        ppr.append(spacing)
        lvl.append(ppr)

        rpr = OxmlElement("w:rPr")
        rfonts = OxmlElement("w:rFonts")
        rfonts.set(qn("w:ascii"), FONT)
        rfonts.set(qn("w:hAnsi"), FONT)
        rpr.append(rfonts)
        size = OxmlElement("w:sz")
        size.set(qn("w:val"), "22")
        rpr.append(size)
        lvl.append(rpr)
        abstract.append(lvl)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(abstract)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int, level: int) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = ensure_child(ppr, "w:numPr")
    ilvl = ensure_child(num_pr, "w:ilvl")
    ilvl.set(qn("w:val"), str(max(0, min(level, 8))))
    num_id_element = ensure_child(num_pr, "w:numId")
    num_id_element.set(qn("w:val"), str(num_id))
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.15


def split_table_row(line: str) -> list[str]:
    body = line.strip()
    if body.startswith("|"):
        body = body[1:]
    if body.endswith("|"):
        body = body[:-1]
    cells = re.split(r"(?<!\\)\|", body)
    return [cell.replace("\\|", "|").strip() for cell in cells]


def is_table_separator(line: str) -> bool:
    cells = split_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells)


def strip_markdown(text: str) -> str:
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    text = text.replace("**", "").replace("*", "").replace("`", "")
    return text


def table_widths(rows: list[list[str]]) -> list[int]:
    header = [strip_markdown(cell).strip().lower() for cell in rows[0]]
    known_layouts = {
        ("sistema ou ferramenta", "papel informado", "estado para o mvp"): [1872, 3089, 4399],
        ("módulo", "incluído no mvp-1"): [2059, 7301],
        ("módulo", "incluído no mvp-2"): [2059, 7301],
        ("entidade ou ação", "fonte recomendada", "regra no vivance"): [2059, 1685, 5616],
        ("ação", "paciente", "médico", "admin técnico"): [3276, 1872, 1872, 2340],
        ("id", "caso de uso", "prioridade", "ator principal"): [936, 4399, 2621, 1404],
    }
    exact = known_layouts.get(tuple(header))
    if exact:
        return exact
    columns = len(rows[0])
    maxima = []
    for column in range(columns):
        values = [len(strip_markdown(row[column])) for row in rows if column < len(row)]
        maxima.append(max(max(values, default=1), 5))
    weights = [math.sqrt(value) for value in maxima]
    minimum = 0.12 if columns >= 4 else 0.16 if columns == 3 else 0.22
    total = sum(weights)
    fractions = [weight / total for weight in weights]
    for _ in range(10):
        below = [idx for idx, fraction in enumerate(fractions) if fraction < minimum]
        if not below:
            break
        fixed = minimum * len(below)
        free = [idx for idx in range(columns) if idx not in below]
        free_total = sum(weights[idx] for idx in free) or 1
        fractions = [
            minimum if idx in below else (1 - fixed) * weights[idx] / free_total
            for idx in range(columns)
        ]
    return column_widths_from_weights(fractions, CONTENT_WIDTH_DXA)


def set_cell_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = ensure_child(tbl_pr, "w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = ensure_child(borders, f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), BORDER)


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    column_count = len(rows[0])
    normalized = [(row + [""] * column_count)[:column_count] for row in rows]
    table = doc.add_table(rows=1, cols=column_count)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.style = "Table Grid"

    for row_index, values in enumerate(normalized):
        row = table.rows[0] if row_index == 0 else table.add_row()
        prevent_row_split(row)
        if row_index == 0:
            repeat_table_header(row)
        for column_index, value in enumerate(values):
            cell = row.cells[column_index]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.style = "Vivance Table Text"
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(strip_markdown(value)) <= 24 and column_index > 0 else WD_ALIGN_PARAGRAPH.LEFT
            add_inline(paragraph, value, bold=row_index == 0)

    widths = table_widths(normalized)
    apply_table_geometry(
        table,
        widths,
        table_width_dxa=CONTENT_WIDTH_DXA,
        indent_dxa=0,
        cell_margins_dxa={"top": 80, "bottom": 80, "start": 120, "end": 120},
    )
    set_cell_borders(table)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    spacer.paragraph_format.line_spacing = 0.7


def load_diagram_font(size: int, bold: bool = False):
    candidates = [
        Path("/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf"),
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size=size)
    return ImageFont.load_default()


def draw_centered_multiline(draw, box, text: str, font, fill: str = "#202124") -> None:
    left, top, right, bottom = box
    words = text.split()
    lines = []
    current = ""
    max_width = right - left - 44
    for word in words:
        candidate = f"{current} {word}".strip()
        width = draw.textbbox((0, 0), candidate, font=font)[2]
        if current and width > max_width:
            lines.append(current)
            current = word
        else:
            current = candidate
    if current:
        lines.append(current)
    line_height = font.size + 8
    total_height = len(lines) * line_height
    y = top + (bottom - top - total_height) / 2
    for line in lines:
        text_box = draw.textbbox((0, 0), line, font=font)
        width = text_box[2] - text_box[0]
        draw.text(((left + right - width) / 2, y), line, font=font, fill=fill)
        y += line_height


def draw_arrow(draw, points, *, color="#5F6368", width=5, head=16) -> None:
    draw.line(points, fill=color, width=width, joint="curve")
    (x1, y1), (x2, y2) = points[-2], points[-1]
    angle = math.atan2(y2 - y1, x2 - x1)
    left = (
        x2 - head * math.cos(angle - math.pi / 6),
        y2 - head * math.sin(angle - math.pi / 6),
    )
    right = (
        x2 - head * math.cos(angle + math.pi / 6),
        y2 - head * math.sin(angle + math.pi / 6),
    )
    draw.polygon([(x2, y2), left, right], fill=color)


def create_mermaid_diagram(labels: dict[str, str], image_path: Path) -> None:
    canvas = Image.new("RGB", (1800, 950), "white")
    draw = ImageDraw.Draw(canvas)
    font = load_diagram_font(34)
    label_font = load_diagram_font(26)
    positions = {
        "A": (70, 130, 500, 300),
        "B": (685, 130, 1115, 300),
        "C": (1300, 130, 1730, 300),
        "D": (1300, 600, 1730, 770),
        "E": (685, 600, 1115, 770),
        "F": (70, 600, 500, 770),
    }
    for node, box in positions.items():
        draw.rounded_rectangle(box, radius=22, fill="#FFFFFF", outline="#BDC1C6", width=4)
        draw_centered_multiline(draw, box, labels.get(node, node), font)

    draw_arrow(draw, [(500, 215), (685, 215)])
    draw_arrow(draw, [(1115, 215), (1300, 215)])
    draw_arrow(draw, [(1515, 300), (1515, 600)])
    draw_arrow(draw, [(1300, 685), (1115, 685)])
    draw_arrow(draw, [(685, 685), (500, 685)])
    draw_arrow(draw, [(70, 685), (25, 685), (25, 215), (685, 215)], color="#80868B", width=4)
    draw.text((45, 430), "Novo contexto", font=label_font, fill="#5F6368")
    draw_arrow(draw, [(1680, 600), (1760, 600), (1760, 300), (1680, 300)], color="#80868B", width=4)
    label = "Corrigir ou rejeitar"
    label_box = draw.textbbox((0, 0), label, font=label_font)
    label_width = label_box[2] - label_box[0]
    draw.text((1740 - label_width, 430), label, font=label_font, fill="#5F6368")
    image_path.parent.mkdir(parents=True, exist_ok=True)
    canvas.save(image_path, format="PNG", optimize=True)


def set_picture_alt_text(paragraph, title: str, description: str) -> None:
    doc_pr = paragraph._p.xpath(".//wp:docPr")
    if doc_pr:
        doc_pr[0].set("title", title)
        doc_pr[0].set("descr", description)


def add_code_block(doc: Document, lines: list[str], language: str, asset_dir: Path) -> None:
    if language.lower() == "mermaid":
        labels = {}
        edges = []
        for line in lines:
            for node, label in re.findall(r"([A-Za-z0-9_]+)\[\"?([^\]]+?)\"?\]", line):
                labels[node] = label.strip('"')
            edge_match = re.match(
                r"\s*([A-Za-z0-9_]+)\s*-->(?:\|\"?([^|]+?)\"?\|)?\s*([A-Za-z0-9_]+)",
                line,
            )
            if edge_match:
                edges.append(edge_match.groups())
        unlabeled = [(a, b) for a, label, b in edges if not label]
        labeled = [(a, label, b) for a, label, b in edges if label]
        if labels and unlabeled:
            image_path = asset_dir / "ciclo_longitudinal_vivance.png"
            create_mermaid_diagram(labels, image_path)
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(4)
            paragraph.paragraph_format.space_after = Pt(2)
            paragraph.add_run().add_picture(str(image_path), width=Inches(6.45))
            set_picture_alt_text(
                paragraph,
                "Ciclo longitudinal de cuidado",
                "Captura, contexto longitudinal, análise e rascunho da IA, revisão e decisão médica, plano ou ação, acompanhamento e retorno ao contexto. Correções retornam à análise.",
            )
            caption = doc.add_paragraph(style="Vivance Metadata")
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption.paragraph_format.left_indent = Inches(0)
            run = caption.add_run("Figura 1 — Ciclo longitudinal de cuidado")
            set_run_font(run, FONT, 9.5, MUTED)
            return

    paragraph = doc.add_paragraph(style="Vivance Code Block")
    set_paragraph_shading(paragraph, CODE_FILL)
    for index, line in enumerate(lines):
        run = paragraph.add_run(line)
        set_run_font(run, MONO_FONT, 9, "202124")
        if index < len(lines) - 1:
            run.add_break()


def add_blockquote(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="Vivance Metadata")
    add_inline(paragraph, text)


def add_body_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="Normal")
    add_inline(paragraph, text)


def convert(source: Path, output: Path) -> dict[str, int]:
    text = source.read_text(encoding="utf-8")
    lines = text.splitlines()
    doc = Document()
    configure_document(doc)

    stats = {"headings": 0, "tables": 0, "list_items": 0, "code_blocks": 0, "paragraphs": 0}
    i = 0
    first_title = True
    active_list_kind = None
    active_list_num_id = None
    paragraph_buffer: list[str] = []

    def flush_paragraph() -> None:
        nonlocal paragraph_buffer
        if paragraph_buffer:
            add_body_paragraph(doc, " ".join(part.strip() for part in paragraph_buffer))
            stats["paragraphs"] += 1
            paragraph_buffer = []

    def reset_list() -> None:
        nonlocal active_list_kind, active_list_num_id
        active_list_kind = None
        active_list_num_id = None

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            flush_paragraph()
            reset_list()
            language = stripped[3:].strip()
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            add_code_block(doc, code_lines, language, output.parent)
            stats["code_blocks"] += 1
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
            flush_paragraph()
            reset_list()
            rows = [split_table_row(line)]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(split_table_row(lines[i]))
                i += 1
            add_markdown_table(doc, rows)
            stats["tables"] += 1
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading_match:
            flush_paragraph()
            reset_list()
            level = len(heading_match.group(1))
            heading_text = heading_match.group(2).strip()
            if first_title and level == 1:
                add_title(doc, heading_text)
                first_title = False
            else:
                mapped_level = min(max(level - 1, 1), 4)
                paragraph = doc.add_paragraph(style=f"Heading {mapped_level}")
                add_inline(paragraph, heading_text)
                stats["headings"] += 1
            i += 1
            continue

        if stripped == "---":
            flush_paragraph()
            reset_list()
            i += 1
            continue

        if stripped.startswith(">"):
            flush_paragraph()
            reset_list()
            quote = stripped[1:].strip()
            add_blockquote(doc, quote)
            stats["paragraphs"] += 1
            i += 1
            continue

        checkbox_match = re.match(r"^(\s*)[-*+]\s+\[([ xX])\]\s+(.+)$", line)
        bullet_match = re.match(r"^(\s*)[-*+]\s+(.+)$", line)
        decimal_match = re.match(r"^(\s*)(\d+)[.)]\s+(.+)$", line)
        list_match = checkbox_match or bullet_match or decimal_match
        if list_match:
            flush_paragraph()
            if checkbox_match:
                kind = "checkbox"
                indent, checked, content = checkbox_match.groups()
                if checked.lower() == "x":
                    content = f"Concluído — {content}"
            elif decimal_match:
                kind = "decimal"
                indent, source_number, content = decimal_match.groups()
            else:
                kind = "bullet"
                indent, content = bullet_match.groups()
            level = min(len(indent.expandtabs(4)) // 2, 8)
            if active_list_kind != kind or active_list_num_id is None:
                active_list_kind = kind
                start_value = int(source_number) if kind == "decimal" else 1
                active_list_num_id = create_numbering(doc, kind, start_value=start_value)
            paragraph = doc.add_paragraph(style="Normal")
            apply_numbering(paragraph, active_list_num_id, level)
            add_inline(paragraph, content)
            stats["list_items"] += 1
            i += 1
            continue

        if not stripped:
            flush_paragraph()
            reset_list()
            i += 1
            continue

        paragraph_buffer.append(stripped)
        i += 1

    flush_paragraph()
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    return stats


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("source", type=Path)
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    stats = convert(args.source, args.output)
    print(f"Created {args.output}")
    print(" ".join(f"{key}={value}" for key, value in stats.items()))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
