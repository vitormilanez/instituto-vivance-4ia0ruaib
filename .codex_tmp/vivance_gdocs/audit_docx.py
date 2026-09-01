#!/usr/bin/env python3
from __future__ import annotations

import re
import sys
import unicodedata
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from lxml import etree


W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
WP = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
NS = {"w": W, "wp": WP, "r": R}


def normalize(text: str) -> str:
    return re.sub(r"\s+", " ", unicodedata.normalize("NFC", text)).strip()


def visible_markdown(text: str) -> str:
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    text = text.replace("**", "").replace("*", "").replace("`", "")
    return normalize(text)


def source_units(path: Path) -> list[str]:
    lines = path.read_text(encoding="utf-8").splitlines()
    units = []
    in_fence = False
    fence_language = ""
    for line in lines:
        stripped = line.strip()
        if stripped.startswith("```"):
            if not in_fence:
                in_fence = True
                fence_language = stripped[3:].strip().lower()
            else:
                in_fence = False
                fence_language = ""
            continue
        if in_fence and fence_language == "mermaid":
            continue
        if not stripped or stripped == "---":
            continue
        if re.fullmatch(r"\|?(?:\s*:?-{3,}:?\s*\|)+", stripped):
            continue
        stripped = re.sub(r"^#{1,6}\s+", "", stripped)
        stripped = re.sub(r"^>\s?", "", stripped)
        stripped = re.sub(r"^[-*+]\s+\[[ xX]\]\s+", "", stripped)
        stripped = re.sub(r"^[-*+]\s+", "", stripped)
        stripped = re.sub(r"^\d+[.)]\s+", "", stripped)
        if stripped.startswith("|"):
            cells = [cell.strip() for cell in re.split(r"(?<!\\)\|", stripped.strip("|"))]
            units.extend(visible_markdown(cell) for cell in cells if visible_markdown(cell))
        else:
            value = visible_markdown(stripped)
            if value:
                units.append(value)
    return units


def paragraph_texts(doc: Document) -> list[str]:
    values = [normalize(paragraph.text) for paragraph in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                values.extend(normalize(paragraph.text) for paragraph in cell.paragraphs)
    return [value for value in values if value]


def numbering_starts(root, numbering_root, texts: list[str]) -> list[int]:
    starts = []
    checks = [
        "O problema prioritário é preparação/documentação médica",
        "Quais perguntas entram na primeira pré-consulta?",
        "O Feegow continuará como prontuário oficial?",
        "Quais dados podem ser enviados a cada fornecedor de IA?",
        "Qual é o nome oficial da marca e o domínio de produção?",
    ]
    for needle in checks:
        paragraph = root.xpath(f".//w:p[contains(string(.), '{needle}')]", namespaces=NS)
        if not paragraph:
            starts.append(-1)
            continue
        num_id = paragraph[0].xpath("string(w:pPr/w:numPr/w:numId/@w:val)", namespaces=NS)
        num = numbering_root.xpath(f".//w:num[@w:numId='{num_id}']", namespaces=NS)
        if not num:
            starts.append(-1)
            continue
        abstract_id = num[0].xpath("string(w:abstractNumId/@w:val)", namespaces=NS)
        start = numbering_root.xpath(
            f"string(.//w:abstractNum[@w:abstractNumId='{abstract_id}']/w:lvl[@w:ilvl='0']/w:start/@w:val)",
            namespaces=NS,
        )
        starts.append(int(start or -1))
    return starts


def main() -> int:
    source = Path(sys.argv[1])
    docx_path = Path(sys.argv[2])
    doc = Document(docx_path)
    failures = []

    heading_counts = {
        name: sum(1 for p in doc.paragraphs if p.style.name == name)
        for name in ("Heading 1", "Heading 2", "Heading 3")
    }
    if heading_counts != {"Heading 1": 26, "Heading 2": 96, "Heading 3": 84}:
        failures.append(f"heading counts: {heading_counts}")

    table_shapes = [(len(table.rows), len(table.columns)) for table in doc.tables]
    expected_shapes = [(8, 3), (14, 2), (6, 2), (12, 3), (12, 4), (17, 4)]
    if table_shapes != expected_shapes:
        failures.append(f"table shapes: {table_shapes}")

    with ZipFile(docx_path) as zf:
        root = etree.fromstring(zf.read("word/document.xml"))
        styles_root = etree.fromstring(zf.read("word/styles.xml"))
        numbering_root = etree.fromstring(zf.read("word/numbering.xml"))
        rels_root = etree.fromstring(zf.read("word/_rels/document.xml.rels"))

    numbered_paragraphs = root.xpath("count(.//w:p[w:pPr/w:numPr])", namespaces=NS)
    if int(numbered_paragraphs) != 1001:
        failures.append(f"numbered paragraphs: {int(numbered_paragraphs)}")

    starts = numbering_starts(root, numbering_root, paragraph_texts(doc))
    if starts != [1, 6, 13, 19, 26]:
        failures.append(f"decision list starts: {starts}")

    title_style = doc.paragraphs[0].style.name
    if title_style == "Title":
        failures.append("title uses Word Title style")
    title_borders = root.xpath(".//w:body/w:p[1]/w:pPr/w:pBdr", namespaces=NS)
    title_underlines = root.xpath(".//w:body/w:p[1]//w:u", namespaces=NS)
    if title_borders or title_underlines:
        failures.append("title contains border or underline")

    table_widths = []
    for table in root.xpath(".//w:tbl", namespaces=NS):
        grid = [int(value) for value in table.xpath("w:tblGrid/w:gridCol/@w:w", namespaces=NS)]
        table_widths.append(grid)
        if sum(grid) != 9360:
            failures.append(f"table width sum: {grid}")
        indent = table.xpath("string(w:tblPr/w:tblInd/@w:w)", namespaces=NS)
        if indent != "0":
            failures.append(f"table indent is {indent}, expected 0")

    expected_widths = [
        [1872, 3089, 4399],
        [2059, 7301],
        [2059, 7301],
        [2059, 1685, 5616],
        [3276, 1872, 1872, 2340],
        [936, 4399, 2621, 1404],
    ]
    if table_widths != expected_widths:
        failures.append(f"table grids: {table_widths}")

    missing_headers = root.xpath("count(.//w:tbl[not(w:tr[1]/w:trPr/w:tblHeader)])", namespaces=NS)
    if int(missing_headers):
        failures.append(f"tables without header flag: {int(missing_headers)}")

    doc_prs = root.xpath(".//wp:docPr", namespaces=NS)
    if len(doc_prs) != 1 or not doc_prs[0].get("descr"):
        failures.append("diagram alt text missing")

    external_links = [
        rel
        for rel in rels_root
        if rel.get("TargetMode") == "External" and "hyperlink" in rel.get("Type", "")
    ]
    if len(external_links) != 5:
        failures.append(f"external hyperlinks: {len(external_links)}")

    all_text = normalize(" ".join(paragraph_texts(doc)))
    missing_units = [unit for unit in source_units(source) if unit not in all_text]
    if missing_units:
        failures.append(f"missing source units ({len(missing_units)}): {missing_units[:12]}")

    print(f"heading_counts={heading_counts}")
    print(f"table_shapes={table_shapes}")
    print(f"numbered_paragraphs={int(numbered_paragraphs)}")
    print(f"decision_list_starts={starts}")
    print(f"external_hyperlinks={len(external_links)}")
    print(f"source_units={len(source_units(source))} missing={len(missing_units)}")
    print(f"diagram_alt_text={doc_prs[0].get('descr') if doc_prs else None}")
    if failures:
        print("FAIL")
        for failure in failures:
            print(f"- {failure}")
        return 1
    print("PASS")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
